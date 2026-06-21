"""Одностраничный проверочный лист UzMRC: доступ + тест-кейсы (Модуль 1/2) + тех. описание."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

GREEN = RGBColor(0x12, 0x7A, 0x41)
GREY = RGBColor(0x55, 0x55, 0x55)
OUT = Path(__file__).resolve().parent.parent / "deliverables" / "UzMRC-Проверочный-лист.docx"


def _set_margins(doc: Document) -> None:
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = Cm(1.2)
    s.left_margin = s.right_margin = Cm(1.5)


def _h(doc: Document, text: str, size: int = 12) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = GREEN


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for k, v in rows:
        c = t.add_row().cells
        c[0].width = Cm(3.5)
        c[1].width = Cm(14)
        rk = c[0].paragraphs[0].add_run(k)
        rk.bold = True
        rk.font.size = Pt(9)
        rv = c[1].paragraphs[0].add_run(v)
        rv.font.size = Pt(9)


def _cases_table(doc: Document, cases: list[tuple[str, str, str]]) -> None:
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, label in enumerate(("№", "Что проверяем / шаги", "Ожидаемый результат")):
        run = hdr[i].paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    widths = (Cm(1.2), Cm(8.3), Cm(8.0))
    for num, step, expect in cases:
        c = t.add_row().cells
        for cell, txt, w in zip(c, (num, step, expect), widths):
            cell.width = w
            run = cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(8.5)


def main() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9)
    _set_margins(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(0)
    r = title.add_run("UzMRC — Нормативный AI-ассистент")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = GREEN
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(4)
    rs = sub.add_run("Проверочный лист: доступ · тест-кейсы (Модуль 1 и 2) · техническое описание")
    rs.font.size = Pt(9)
    rs.font.color.rgb = GREY

    _h(doc, "1. Доступ к системе")
    _kv_table(doc, [
        ("URL", "https://89.167.15.225.sslip.io"),
        ("Логин", "admin@uzmrc.io"),
        ("Пароль", "______________________  (выдаётся администратором; по запросу заводим отдельный проверочный аккаунт)"),
        ("Браузер", "Chrome / Edge / Safari, актуальная версия"),
    ])

    _h(doc, "2. Модуль 1 — Чат с цитированием (вопрос → ответ со ссылками на нормы)")
    _cases_table(doc, [
        ("1.1", "Открыть базу → «Чат». Спросить: «Какова антикоррупционная политика компании?»",
         "Структурированный ответ; в тексте зелёные цитаты 【1】…; показана «уверенность, %»."),
        ("1.2", "Кликнуть по цитате 【1】 в ответе.",
         "Всплывает фрагмент-источник: имя файла, страница, точный текст; кнопка «Открыть документ»."),
        ("1.3", "Задать вопрос на узбекском: «kuzatuv kengashi vakolatlari».",
         "Ответ по положению о набсовете; кросс-язык работает (uz-запрос → uz/ru источники)."),
        ("1.4", "Спросить заведомо отсутствующее в базе (напр. «график отпусков на 2027»).",
         "Честный ответ «в документах не покрыто», без выдуманных фактов (анти-галлюцинация)."),
        ("1.5", "Проверить блок «Источники» под ответом.",
         "Перечень источников с нумерацией [n], каждый кликабелен → превью фрагмента."),
    ])

    _h(doc, "3. Модуль 2 — Сверка документов (новый приказ ↔ действующие нормы)")
    _cases_table(doc, [
        ("2.1", "Открыть базу → «Сравнить документ». Загрузить приказ (.pdf/.txt/.md/.xlsx).",
         "Идёт прогресс (загрузка + анализ); по завершении — отчёт с KPI-карточками."),
        ("2.2", "Пункт приказа, ослабляющий/нарушающий норму (напр. «подарки без ограничений»).",
         "Вердикт «Противоречие» + цитата нормы + обоснование + рекомендация «отклонить»."),
        ("2.3", "Пункт, совпадающий с действующей нормой.",
         "Вердикт «Дубль»; рекомендация «принять»; цитата подтверждена (grounding)."),
        ("2.4", "Пункт, которого нет в нормах базы (новая тема).",
         "Вердикт «Пробел» / «Дополнение»; рекомендация «согласовать»."),
        ("2.5", "Нажать фильтр-чипы (Противоречия / Дубли / …) и «Скачать .md» / «Печать».",
         "Список фильтруется по типу; отчёт выгружается в .md и печатается в PDF."),
    ])

    _h(doc, "4. Техническое описание (кратко)")
    tech = doc.add_paragraph()
    tech.paragraph_format.space_after = Pt(0)
    lines = [
        ("Назначение: ", "AI-ассистент по нормативной базе УзКРИ — поиск с цитированием (Модуль 1) и сверка новых приказов с действующими нормами (Модуль 2)."),
        ("Корпус: ", "≈500 документов сайта uzmrc.uz (PDF + страницы), языки ru/uz/en, кросс-язычный поиск."),
        ("Архитектура: ", "Фронт — Next.js 16 / React 19; бэкенд — FastAPI; хранилище — PostgreSQL + pgvector; деплой — Docker + Caddy (HTTPS)."),
        ("Модели: ", "эмбеддинги Voyage voyage-3.5 (1024-мерные); ответы/судья — Cerebras gpt-oss-120b; реранкер — Voyage rerank-2.5; OCR сканов — qwen-vl."),
        ("Достоверность: ", "каждая цитата дословно сверяется с фрагментом-источником (grounding); «уверенность» ограничена долей подтверждённых цитат — система не приписывает нормам то, чего в них нет."),
        ("Безопасность: ", "доступ по логину (JWT), HTTPS, защита роутов; документ при сверке НЕ добавляется в базу."),
    ]
    for label, text in lines:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(0)
        rl = p.add_run(label)
        rl.bold = True
        rl.font.size = Pt(8.5)
        rt = p.add_run(text)
        rt.font.size = Pt(8.5)

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(6)
    rf = foot.add_run("UzMRC · нормативный AI-ассистент · версия для проверки · 2026")
    rf.font.size = Pt(7.5)
    rf.font.color.rgb = GREY

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("saved:", OUT)


if __name__ == "__main__":
    main()
