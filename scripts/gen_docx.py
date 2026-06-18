"""Генерация DOCX технического описания UzMRC."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "deliverables", "UzMRC-Техническое-описание.docx")
os.makedirs(os.path.dirname(OUT), exist_ok=True)

ACCENT = RGBColor(0x5B, 0x3D, 0xF5)  # фиолетовый бренда
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GREY = RGBColor(0x66, 0x66, 0x66)

doc = Document()
styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p

def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs:
        r.font.color.rgb = DARK
    return p

def para(t, italic=False, color=None, size=None, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.italic = italic
    r.bold = bold
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return p

def bullets(items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, hh in enumerate(headers):
        hdr[i].paragraphs[0].add_run(hh).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            cells[i].text = str(c)
    doc.add_paragraph()
    return t

# --- Титул ---
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("UzMRC")
r.font.size = Pt(40)
r.font.bold = True
r.font.color.rgb = ACCENT
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("ИИ-ассистент по нормативным документам")
rs.font.size = Pt(16)
rs.font.color.rgb = DARK
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
rm = meta.add_run("Полное техническое описание · MVP · 18 июня 2026")
rm.font.size = Pt(11)
rm.font.color.rgb = GREY
doc.add_paragraph()

# --- 1 ---
h1("1. Назначение")
para("UzMRC — система для работы с внутренними нормативными документами компании. Решает две задачи:")
bullets([
    "Чат по документам — ответы на вопросы (рус/узб) со ссылкой на источник: файл, страница, точная цитата. Кросс-языковой поиск. Защита от галлюцинаций.",
    "Сравнение документов — разбор нового приказа/регламента на пункты и сверка каждого с действующей нормативной базой: противоречия, дубли, пробелы, дополнения. Отчёт с обоснованием, цитатой нормы и рекомендацией.",
])
para("Система мультитенантная по принципу «один RAG = одна изолированная база»: своя Qdrant-коллекция, свой каталог файлов, свои чанки в Postgres. Удаление RAG каскадно убирает всё.")

# --- 2 ---
h1("2. Архитектура и стек")
table(["Слой", "Технология"], [
    ["Frontend", "React 18 + Vite + TypeScript + Tailwind (SPA), nginx"],
    ["Backend", "FastAPI (Python 3.12), async SQLAlchemy, Pydantic v2"],
    ["Реляционная БД", "PostgreSQL 16 (метаданные, чанки, FTS, кэш эмбеддингов)"],
    ["Векторная БД", "Qdrant (плотные векторы, cosine)"],
    ["Эмбеддинги", "Voyage voyage-3.5 (1024-dim, multilingual)"],
    ["LLM-судья / rerank", "Cerebras gpt-oss-120b (free 1M ток/день), фолбэк OpenRouter"],
    ["Оркестрация", "Docker Compose (4 контейнера)"],
    ["Миграции", "Alembic"],
    ["Авторизация", "JWT (HS256), bcrypt, закрытая регистрация"],
])

# --- 3 ---
h1("3. Модуль 1 — Чат по документам (агентный RAG)")
h2("3.1 Гибридный поиск")
bullets([
    "Dense — эмбеддинг запроса Voyage → поиск в Qdrant (cosine).",
    "Sparse — Postgres full-text search (ts_rank_cd), язык настраивается на RAG (russian/simple).",
    "Fusion — Reciprocal Rank Fusion (RRF, k=60), затем топ-K судье/агенту.",
    "Rerank — LLM-rerank пула кандидатов для повышения точности.",
])
h2("3.2 Агент (ReAct + структурированный вывод)")
bullets([
    "Цикл рассуждение→действие→наблюдение, до 40 шагов.",
    "На каждом шаге LLM возвращает строго один JSON-объект, валидируемый Pydantic (структурированная генерация, SGR).",
    "Query Router до цикла: regex-эвристики + LLM-fallback при низкой уверенности.",
    "Инструменты: hybrid_search, dense/sparse_search, decompose_and_search, hyde_search, exact_lookup, fetch_page, fetch_document, list_files, cache_fact/recall_fact, rerank_pool.",
    "Финал/эскалация — отдельные виды шага: final (с цитатами) или escalate (нет ответа в базе).",
])
h2("3.3 Grounding-проверка (анти-галлюцинации)")
para("Каждая цитата финального ответа проверяется на наличие в указанном чанке: подстрока на нормализованном тексте → агрессивная нормализация (переносы/пунктуация/гомоглифы OCR) → fuzzy-сопоставление (порог ≥ 0.78).")
para("Уверенность ответа ограничивается долей подтверждённых цитат. Если в базе нет нормы — система не отвечает и не фабрикует ссылку.", bold=True)

# --- 4 ---
h1("4. Модуль 2 — Сравнение документов")
para("Документ разбивается на пункты; по каждому — hybrid-поиск + rerank (пул 10 → топ-5) → LLM-судья классифицирует отношение пункта к ближайшей норме базы.")
table(["Тип", "Значение"], [
    ["Противоречие (conflict)", "пункт нельзя исполнить, не нарушив действующую норму"],
    ["Дубль (duplicate)", "повторяет существующую норму"],
    ["Дополнение (addition)", "расширяет/уточняет, не противореча"],
    ["Пробел (gap)", "в базе нет нормы по теме"],
])
para("По каждой находке: обоснование, цитата нормы (с проверкой grounding), рекомендация, уверенность.")
para("Асинхронное исполнение: POST /compare возвращает 202 + run_id, не блокируя запрос; прогресс через SSE. Документ на 10 пунктов — ~20–25 с. Лимит — 120 пунктов за прогон.")

# --- 5 ---
h1("5. Оптимизации и надёжность")
bullets([
    "Кэш эмбеддингов (content-addressed по sha256(provider+model+text)): повторная индексация не жжёт квоту — реиндекс из ~6 мин → ~5 с (×77).",
    "Token-aware суб-батчинг эмбеддингов — один запрос не пробивает TPM-лимит тарифа.",
    "OCR-fallback (Vision-LLM) для PDF со сканами при <200 извлечённых символов.",
    "Иерархия фолбэков LLM: Cerebras → OpenRouter → Gemini (резерв эмбеддингов).",
])

# --- 6 ---
h1("6. Безопасность")
bullets([
    "JWT startup-guard, закрытая регистрация, создание пользователей только админом.",
    "Rate-limit логина (10/мин/IP), отключаемые /docs, nginx security-заголовки (CSP/HSTS).",
    "Per-RAG изоляция данных на всех слоях (FS/Postgres/Qdrant/API).",
    "Секреты только в .env (вне репозитория). Бэкап БД — scripts/backup-db.sh.",
])

# --- 7 ---
h1("7. База знаний (текущий корпус MVP)")
table(["Показатель", "Значение"], [
    ["Документов (уникальных)", "50 (.txt из PDF сайта uzmrc.uz)"],
    ["Фрагментов (чанков)", "1 630"],
    ["Чанков на документ (среднее)", "32.6"],
    ["Токенов в индексе", "643 974"],
    ["Модель эмбеддингов", "voyage-3.5, 1024-dim"],
    ["Источник", "≈35 нормативных актов + ≈12 аналитических обзоров (≈595 стр.)"],
    ["Статус", "Готова (ready)"],
])
para("Цифры доступны в приложении на странице «О системе» (тянутся вживую из GET /api/rags/{id}/stats). Это демо-набор, а не полная база.", italic=True)

# --- 8 ---
h1("8. Что входит в MVP / ограничения")
para("Входит:", bold=True)
bullets([
    "Чат по базе со ссылками на источник.",
    "Сравнение документа с нормами.",
    "Кросс-язык ru↔uz.",
    "Экспорт отчёта (.md/PDF).",
])
para("Не входит:", bold=True)
bullets([
    "Анализ портфеля, маскирование персональных данных, генерация презентаций, парсинг рынка/OLX.",
])
para("Ограничения:", bold=True)
bullets([
    "До 120 пунктов за прогон сравнения.",
    "Загрузка новых документов требует интернета и платных ключей Voyage.",
    "Поиск/сравнение по уже загруженной базе — быстрые. Запуск локально через Docker.",
])

# --- 9 ---
h1("9. Развёртывание")
para("git clone https://github.com/Yersultan04/uzmrc.git → cd uzmrc/rag-cms → положить .env → "
     "docker compose -f docker-compose.prod.yml up -d --build → UI на http://localhost:8090.")
para("Для постоянного публичного адреса — деплой на VPS (домен + reverse-proxy); для временного — Cloudflare quick tunnel (бесплатно, URL эфемерный). См. notes/DEPLOY-RUNBOOK.md и QUICKSTART.md.")

# --- 10 ---
h1("10. Качество (метрики)")
bullets([
    "Кросс-язык: запрос «антикоррупционная политика» (рус) находит русский antokorrup-politika (0.80) и узбекские korrupsiyagaqarshi (0.79).",
    "Демо-приказ сравнения: 10 пунктов → 4 противоречия / 1 дубль / 1 дополнение / 4 пробела, цитаты норм 5/7, ~23 с.",
    "Eval-наборы (retrieval_gold.jsonl, compare_gold.jsonl) + baseline в репозитории.",
    "10 реальных примеров работы — см. notes/10-EXAMPLES.md.",
])

doc.save(OUT)
print("OK ->", OUT)
